"""Build the current Barber Hub guide as a polished Word handbook.

Preset: compact_reference_guide.
Named visual override: Barber Hub brand palette (warm gold/dark brown) replaces
the default blue accents while preserving the preset geometry and rhythm.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "GUIA_COMPLETO_BARBER_HUB_1_10_0.docx"
LOGO = ROOT / "img" / "branding" / "barber-hub-horizontal-fundo-claro.png"

SOURCES = [
    (ROOT / "docs" / "GUIA_COMPLETO_DO_PROJETO.md", None),
    (ROOT / "docs" / "PESQUISA_VALIDACAO_BARBER_BEAUTY_HUB.md", "Apêndice A — Pesquisa e validação"),
    (ROOT / "docs" / "RELATORIO_SEGURANCA_1_10.md", "Apêndice B — Segurança"),
    (ROOT / "docs" / "MIGRATIONS_DEPLOY_1_10.md", "Apêndice C — Banco e deploy"),
    (ROOT / "docs" / "HOMOLOGACAO_FINAL_1_10.md", "Apêndice D — Homologação"),
]

INK = "231F1A"
GOLD = "A97715"
GOLD_DARK = "6B4E0F"
MUTED = "6B655C"
LINE = "D8CCB7"
PALE = "F4EFE6"
PALE_GOLD = "FBF6E9"
WHITE = "FFFFFF"
PAGE_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_repeatable_list_geometry(paragraph) -> None:
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


def add_inline(paragraph, text: str, *, color: str | None = None) -> None:
    text = text.replace("&#x20;", " ").strip()
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]+\))")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            if color:
                run.font.color.rgb = rgb(color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = rgb(GOLD_DARK)
        else:
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            run = paragraph.add_run(f"{label} ({url})")
            run.font.color.rgb = rgb(GOLD_DARK)
            run.underline = True
        if color and token.startswith("**"):
            run.font.color.rgb = rgb(color)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        if color:
            run.font.color.rgb = rgb(color)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr_text, fld_char_2])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    headings = {
        "Heading 1": (16, GOLD_DARK, 18, 10),
        "Heading 2": (13, GOLD, 14, 7),
        "Heading 3": (12, GOLD_DARK, 10, 5),
    }
    for name, (size, color, before, after) in headings.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.color.rgb = rgb(INK)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    code = styles.add_style("BH Code", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = "Consolas"
    code.font.size = Pt(8.5)
    code.font.color.rgb = rgb(INK)
    code.paragraph_format.left_indent = Inches(0.16)
    code.paragraph_format.right_indent = Inches(0.16)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.0

    callout = styles.add_style("BH Callout", WD_STYLE_TYPE.PARAGRAPH)
    callout.font.name = "Calibri"
    callout.font.size = Pt(10.5)
    callout.font.color.rgb = rgb(INK)
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.18)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(8)
    callout.paragraph_format.line_spacing = 1.2

    for sec in doc.sections:
        header = sec.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("BARBER HUB  ·  GUIA DO PRODUTO E DA OPERAÇÃO")
        run.font.name = "Calibri"
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.color.rgb = rgb(MUTED)
        footer = sec.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp.paragraph_format.space_before = Pt(0)
        fr = fp.add_run("The Gamers Tech  ·  1.10.0  ·  ")
        fr.font.name = "Calibri"
        fr.font.size = Pt(8.5)
        fr.font.color.rgb = rgb(MUTED)
        add_field(fp, "PAGE")


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(30)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        picture = p.add_run().add_picture(str(LOGO), width=Inches(3.1))
        picture._inline.docPr.set("descr", "Logotipo horizontal do Barber Hub")
        picture._inline.docPr.set("title", "Barber Hub")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(48)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("GUIA COMPLETO")
    run.font.name = "Calibri"
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = rgb(GOLD_DARK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Barber Hub 1.10.0")
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = rgb(INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Produto · arquitetura · segurança · operação · deploy · pesquisa")
    run.font.name = "Calibri"
    run.font.size = Pt(11.5)
    run.font.color.rgb = rgb(MUTED)

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    values = [
        ("Produto", "Marketplace e gestão para barbearias"),
        ("Situação", "Código em homologação; migrations 29–31 pendentes"),
        ("Atualização", "2 de setembro de 2026"),
    ]
    for r_idx, (label, value) in enumerate(values):
        table.cell(r_idx, 0).text = label
        table.cell(r_idx, 1).text = value
        set_cell_shading(table.cell(r_idx, 0), PALE)
        table.cell(r_idx, 0).paragraphs[0].runs[0].bold = True
        keep_row_together(table.rows[r_idx])
    repeat_table_header(table.rows[0])
    set_table_geometry(table, [2100, 7260])

    p = doc.add_paragraph(style="BH Callout")
    p.paragraph_format.space_before = Pt(22)
    add_inline(p, "Documento canônico da versão 1.10. O código ainda não autoriza publicação: banco, serviços externos e homologação precisam ser concluídos.")
    shade_paragraph(p, PALE_GOLD)
    doc.add_page_break()


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    if cols == 2:
        widths = [2700, 6660]
    elif cols == 3:
        widths = [2300, 4100, 2960]
    elif cols == 4:
        widths = [1700, 3100, 2800, 1760]
    else:
        base = PAGE_DXA // cols
        widths = [base] * cols
        widths[-1] += PAGE_DXA - sum(widths)

    for r_idx, row in enumerate(rows):
        keep_row_together(table.rows[r_idx])
        for c_idx in range(cols):
            text = row[c_idx].strip() if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            add_inline(p, text)
            for run in p.runs:
                run.font.size = Pt(9.3)
            if r_idx == 0:
                set_cell_shading(cell, PALE)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = rgb(GOLD_DARK)
    repeat_table_header(table.rows[0])
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def parse_markdown(doc: Document, path: Path, part_title: str | None) -> None:
    if part_title:
        doc.add_page_break()
        p = doc.add_paragraph(part_title, style="Heading 1")
        p.paragraph_format.space_before = Pt(0)

    lines = path.read_text(encoding="utf-8").splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []
    first_h1_skipped = False

    while index < len(lines):
        raw = lines[index].rstrip()
        stripped = raw.strip()

        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph(style="BH Code")
                p.add_run("\n".join(code_lines))
                shade_paragraph(p, PALE)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue

        if not stripped or stripped == "---":
            index += 1
            continue

        if stripped.startswith("|") and index + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-{3,}", lines[index + 1]):
            rows: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                values = [item.strip() for item in lines[index].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", item.replace(" ", "")) for item in values):
                    rows.append(values)
                index += 1
            add_table(doc, rows)
            continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1 and not first_h1_skipped:
                first_h1_skipped = True
                index += 1
                continue
            style = "Heading 1" if level <= 2 else "Heading 2" if level == 3 else "Heading 3"
            p = doc.add_paragraph(style=style)
            add_inline(p, text)
            index += 1
            continue

        if stripped.startswith(">"):
            p = doc.add_paragraph(style="BH Callout")
            add_inline(p, stripped.lstrip("> "))
            shade_paragraph(p, PALE_GOLD)
            index += 1
            continue

        if re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            set_repeatable_list_geometry(p)
            add_inline(p, re.sub(r"^[-*]\s+", "", stripped))
            index += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            set_repeatable_list_geometry(p)
            add_inline(p, re.sub(r"^\d+\.\s+", "", stripped))
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            nxt = lines[index].strip()
            if (
                not nxt
                or nxt.startswith(("#", "```", ">", "|"))
                or re.match(r"^[-*]\s+", nxt)
                or re.match(r"^\d+\.\s+", nxt)
                or nxt == "---"
            ):
                break
            paragraph_lines.append(nxt)
            index += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(paragraph_lines))


def finalize(doc: Document) -> None:
    props = doc.core_properties
    props.title = "Guia Completo Barber Hub 1.10.0"
    props.subject = "Produto, arquitetura, segurança, operação, deploy e validação"
    props.author = "The Gamers Tech"
    props.keywords = "Barber Hub, Beauty Hub, FastAPI, Supabase, PWA, segurança"
    props.comments = "Atualizado para a versão 1.10.0"
    doc.settings.element.append(OxmlElement("w:updateFields"))
    doc.save(OUTPUT)


def main() -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    for path, part_title in SOURCES:
        parse_markdown(doc, path, part_title)
    finalize(doc)
    print(OUTPUT)


if __name__ == "__main__":
    main()
